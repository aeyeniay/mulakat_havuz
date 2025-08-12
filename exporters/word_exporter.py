"""
WORD BELGESİ EXPORT SİSTEMİ
===========================

Üretilen soruları profesyonel Word belgesine çeviren sistem.
"""

import logging
from datetime import datetime
from typing import Dict, Any, List
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    raise ImportError("python-docx kütüphanesi yüklü değil. 'pip install python-docx' komutu ile yükleyin.")

from config.rubric_system import DIFFICULTY_LABELS
from utils.file_helpers import FileHelper

logger = logging.getLogger(__name__)

class WordExporter:
    """Word belgesi export sınıfı"""
    
    def __init__(self):
        """Word exporter başlatıcı"""
        self.document = None
    
    def create_document(self) -> Document:
        """Yeni Word belgesi oluştur"""
        self.document = Document()
        self._setup_document_styles()
        return self.document
    
    def _setup_document_styles(self):
        """Belge stillerini ayarla"""
        # Başlık stilleri
        title_style = self.document.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Alt başlık stilleri
        subtitle_style = self.document.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.name = 'Arial'
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.bold = True
        subtitle_style.paragraph_format.space_before = Pt(12)
        subtitle_style.paragraph_format.space_after = Pt(6)
        
        # Soru stilleri
        question_style = self.document.styles.add_style('Custom Question', WD_STYLE_TYPE.PARAGRAPH)
        question_style.font.name = 'Arial'
        question_style.font.size = Pt(11)
        question_style.font.bold = True
        question_style.paragraph_format.space_before = Pt(8)
        question_style.paragraph_format.space_after = Pt(4)
        
        # Cevap stilleri
        answer_style = self.document.styles.add_style('Custom Answer', WD_STYLE_TYPE.PARAGRAPH)
        answer_style.font.name = 'Arial'
        answer_style.font.size = Pt(10)
        answer_style.paragraph_format.space_after = Pt(8)
        answer_style.paragraph_format.left_indent = Inches(0.25)

        # Kod stilleri (tek satır veya çok satır kod blokları için)
        code_style = self.document.styles.add_style('Custom Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.font.bold = False
        # 3 boşluk kadar görsel girinti yaklaşığı: küçük bir sol iç boşluk
        code_style.paragraph_format.left_indent = Inches(0.2)
        code_style.paragraph_format.space_after = Pt(0)
    
    def export_questions(
        self,
        questions_data: Dict[str, Any],
        job_description: str,
        output_path: str
    ) -> bool:
        """
        Soruları Word belgesine export et.
        
        Args:
            questions_data (dict): Soru verileri
            job_description (str): İlan metni
            output_path (str): Çıktı dosyası yolu
            
        Returns:
            bool: Başarı durumu
        """
        try:
            # Yeni belge oluştur
            self.create_document()
            
            # Belge başlığını ekle
            self._add_document_header(questions_data)
            
            # Soruları ekle
            self._add_questions_sections(questions_data)
            
            # Dosyayı kaydet
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            self.document.save(str(output_file))
            logger.info(f"Word belgesi başarıyla kaydedildi: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Word export hatası: {e}")
            return False
    
    def _add_document_header(self, questions_data: Dict[str, Any]):
        """Belge başlığını ekle"""
        role = questions_data.get("role", "Bilinmeyen Pozisyon")
        salary_coefficient = questions_data.get("salary_coefficient", 2)
        
        # Ana başlık - katsayı bilgisiyle
        title = f"{role.upper()} {salary_coefficient}x MÜLAKAT SORULARI"
        title_para = self.document.add_paragraph(title, style='Custom Title')
        
        # Tarih
        date_para = self.document.add_paragraph(f"Oluşturulma Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()  # Boş satır
    
    def _add_job_description_section(self, job_description: str):
        """İlan bilgileri bölümünü ekle"""
        self.document.add_paragraph("İlan Bilgileri", style='Custom Subtitle')
        self.document.add_paragraph("-" * 50)
        
        # İlan metnini paragraf paragraf ekle
        for paragraph in job_description.split('\n'):
            if paragraph.strip():
                self.document.add_paragraph(paragraph.strip())
        
        self.document.add_paragraph()  # Boş satır
    
    def _add_rubric_distribution_section(self, questions_data: Dict[str, Any]):
        """Rübrik dağılımı bölümünü ekle"""
        salary_coefficient = questions_data.get("salary_coefficient", 2)
        
        # İlk sorudan difficulty_distribution'ı al
        difficulty_distribution = None
        for category_questions in questions_data.get("questions", {}).values():
            if category_questions and len(category_questions) > 0:
                difficulty_distribution = category_questions[0].get("difficulty_distribution")
                break
        
        if not difficulty_distribution:
            return
        
        difficulty_label = DIFFICULTY_LABELS.get(salary_coefficient, {"label": f"{salary_coefficient}x"})
        
        self.document.add_paragraph(f"Rübrik Dağılımı ({difficulty_label['label']} Seviye)", style='Custom Subtitle')
        self.document.add_paragraph("-" * 50)
        
        # Rübrik seviyelerini listele
        rubric_mapping = {
            "K1_Temel_Bilgi": "K1 - Temel Bilgi",
            "K2_Uygulamali": "K2 - Uygulamalı",
            "K3_Hata_Cozumleme": "K3 - Hata Çözümleme",
            "K4_Tasarim": "K4 - Tasarım",
            "K5_Stratejik": "K5 - Stratejik"
        }
        
        for level, percentage in difficulty_distribution.items():
            level_name = rubric_mapping.get(level, level)
            self.document.add_paragraph(f"• {level_name}: %{percentage}")
        
        self.document.add_paragraph()  # Boş satır
    
    def _add_questions_sections(self, questions_data: Dict[str, Any]):
        """Soru bölümlerini ekle"""
        questions = questions_data.get("questions", {})
        
        category_names = {
            "professional_experience": "Mesleki Deneyim Soruları",
            "theoretical_knowledge": "Teorik Bilgi Soruları",
            "practical_application": "Pratik Uygulama Soruları"
        }
        
        for category_code, category_questions in questions.items():
            if not category_questions:
                continue
            
            category_name = category_names.get(category_code, category_code.replace("_", " ").title())
            question_count = len([q for q in category_questions if q.get("success", False)])
            
            # Kategori başlığı
            section_title = f"{category_name} ({question_count} Soru)"
            self.document.add_paragraph(section_title, style='Custom Subtitle')
            self.document.add_paragraph("-" * len(section_title))
            
            # Her soruyu ekle
            for i, question_data in enumerate(category_questions, 1):
                if not question_data.get("success", False):
                    continue
                
                self._add_single_question(category_code, i, question_data)
            
            self.document.add_paragraph()  # Kategori arası boşluk
    
    def _add_single_question(self, category_code: str, question_number: int, question_data: Dict[str, Any]):
        """Tek bir soruyu ekle"""
        question_text = question_data.get("question", "Soru metni eksik")
        expected_answer = question_data.get("expected_answer", "Beklenen cevap eksik")

        # Soru metni ile olası kod bloğunu ayır
        clean_question, code_block = self._split_question_and_code(question_text)

        # Soru başlığı ve metni (bold)
        self.document.add_paragraph(f"{question_number}. {clean_question}", style='Custom Question')

        # Kod bloğu sadece 'practical_application' kategorisinde yazılsın
        if code_block and category_code == 'practical_application':
            # Görselleştirme öncesi kodu normalize et (kaçışları düzelt, eksik kapatmaları tamamla)
            code_block = self._normalize_code_block_for_display(code_block)
            self.document.add_paragraph("")  # Boş satır
            for line in code_block.splitlines():
                code_para = self.document.add_paragraph(style='Custom Code')
                run = code_para.add_run(f"   {line.rstrip()}")
                run.bold = False
            # Kod ile cevap arasında bir boş satır
            self.document.add_paragraph("")

        # Beklenen cevap
        if expected_answer:
            self.document.add_paragraph(f"Beklenen Cevap: {expected_answer}", style='Custom Answer')

        self.document.add_paragraph()  # Soru arası boşluk

    def _split_question_and_code(self, question_text: str) -> tuple:
        """Soru metninden olası kod bloğunu ayır.

        Varsayım: Kod sorularında, soru cümlesini ilk satır oluşturur ve
        bunu takip eden satırlar kod satırlarıdır. Kod, markdown blokları
        olmadan düz metindir ve satırlar \n ile ayrılır.

        Returns:
            (clean_question, code_block)  
            code_block None ise kod bulunmamıştır.
        """
        if not question_text:
            return "Soru metni eksik", None

        lines = [ln for ln in question_text.splitlines()]
        if len(lines) <= 1:
            return question_text.strip(), None

        first = lines[0].strip()
        candidate = "\n".join(lines[1:]).strip()

        if self._looks_like_code_block(candidate):
            return first, candidate

        # Kod gibi görünmüyor; tüm metni tek soru kabul et
        return question_text.strip(), None

    def _looks_like_code_block(self, text: str) -> bool:
        """Basit sezgisel kontrollerle bir metnin kod olup olmadığını tahmin et."""
        if not text:
            return False

        code_indicators = [
            # Yaygın anahtar kelimeler
            r"^\s*(for|if|while|def|class|try|catch|public|private|protected|static|using|import|from)\b",
            r"\bConsole\.",
            r"\bprint\(",
            r"\bSystem\.",
            r"\bvar\s+\w+\s*=",
            r"\blet\s+\w+\s*=",
            r"\bconst\s+\w+\s*=",
            r"\bNew-[A-Za-z]+\b",   # PowerShell
            r"\bGet-[A-Za-z]+\b",
            r"\bSet-[A-Za-z]+\b",
            r"SELECT\s+.+\s+FROM",
            r"CREATE\s+TABLE",
            r"INSERT\s+INTO",
            r"UPDATE\s+\w+\s+SET",
            r"DELETE\s+FROM",
            # Semboller
            r"[{};=()<>\[\]]",
            r":\s*$",  # Python bloğu
        ]

        import re
        lines = [ln for ln in text.splitlines() if ln.strip()]
        if not lines:
            return False

        # Eğer iki veya daha fazla satır varsa ve en az bir satırda belirgin kod göstergesi varsa
        indicator_regexes = [re.compile(pat, re.IGNORECASE) for pat in code_indicators]
        hits = 0
        for ln in lines:
            if any(r.search(ln) for r in indicator_regexes):
                hits += 1
        return hits >= 1 and len(lines) >= 1

    def _normalize_code_block_for_display(self, text: str) -> str:
        """Kod bloğunu Word çıktısı için güvenli ve mümkün olduğunca tam hale getirir.

        - Yaygın kaçışları çözer (\\\", \\')
        - \n literallerini gerçek satır sonlarına çevirir
        - Parantez/ayraç ve çift tırnakları dengelemeye çalışır
        - Console.WriteLine( … ) gibi yaygın kalıpları kapatır
        """
        import re
        if not text:
            return text
        # Kaçışları çöz
        normalized = text.replace('\\"', '"').replace("\\'", "'")
        # \n literalini gerçek yeni satıra çevir
        normalized = normalized.replace('\\n', '\n')

        # Satır bazlı düzeltmeler
        lines = [ln.rstrip() for ln in normalized.splitlines()]
        fixed_lines = []
        for ln in lines:
            # Console.WriteLine veya benzeri açık parantezle kalan kalıplar
            if re.search(r"Console\.Write(Line|)\s*\(.*$", ln) and not re.search(r"\)\s*;\s*$", ln):
                # Kapanışı ekle
                if ln.endswith('(') or ln.endswith('( '):
                    ln = ln + ')'
                if not ln.endswith(')'):
                    # Eksik ) varsa ekle
                    open_paren = ln.count('(') - ln.count(')')
                    ln = ln + (')' * max(open_paren, 1))
                if not ln.strip().endswith(';'):
                    ln = ln + ';'
            # Satır içinde tek sayıda çift tırnak varsa kapat
            if ln.count('"') % 2 == 1:
                ln = ln + '"'
            fixed_lines.append(ln)

        normalized = '\n'.join(fixed_lines)

        # Blok seviyesinde ayraç dengeleme: (), {}, []
        def balance(block: str, open_ch: str, close_ch: str) -> str:
            count = block.count(open_ch) - block.count(close_ch)
            if count > 0:
                return block + (close_ch * count)
            return block

        normalized = balance(normalized, '(', ')')
        normalized = balance(normalized, '{', '}')
        normalized = balance(normalized, '[', ']')
        return normalized

    # Dil tespiti ve gösterimi kullanıcı talebiyle kaldırıldı
    
    def _add_document_footer(self):
        """Belge alt bilgisini ekle"""
        self.document.add_paragraph()
        self.document.add_paragraph("-" * 80)
        
        footer_para = self.document.add_paragraph("Bu belge Mülakat Soru Havuzu sistemi tarafından otomatik olarak üretilmiştir.")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.italic = True
    
    def generate_filename(
        self,
        role_name: str,
        salary_coefficient: int,
        base_dir: str = "data/word_exports"
    ) -> str:
        """
        Dosya ismi oluştur.
        
        Args:
            role_name (str): Pozisyon ismi
            salary_coefficient (int): Maaş katsayısı
            base_dir (str): Ana dizin
            
        Returns:
            str: Tam dosya yolu
        """
        # Güvenli dosya ismi oluştur - katsayı bilgisiyle
        safe_role_name = FileHelper.get_safe_filename(role_name)
        
        filename = f"{safe_role_name}_{salary_coefficient}x.docx"
        
        return str(Path(base_dir) / filename)
    
    def export_multiple_roles(
        self,
        multiple_questions_data: List[Dict[str, Any]],
        job_descriptions: Dict[str, str],
        output_dir: str = "data/word_exports"
    ) -> Dict[str, Any]:
        """
        Birden fazla rol için Word belgeleri oluştur.
        
        Args:
            multiple_questions_data (list): Çoklu soru verileri
            job_descriptions (dict): İlan metinleri
            output_dir (str): Çıktı dizini
            
        Returns:
            dict: Export sonuçları
        """
        results = {
            "success": True,
            "exported_files": [],
            "failed_exports": []
        }
        
        for questions_data in multiple_questions_data:
            try:
                role = questions_data.get("role", "unknown")
                salary_coefficient = questions_data.get("salary_coefficient", 2)
                
                # Dosya yolu oluştur
                output_path = self.generate_filename(role, salary_coefficient, output_dir)
                
                # İlan metnini al
                job_description = job_descriptions.get(role, f"{role} pozisyonu için iş tanımı")
                
                # Export et
                export_success = self.export_questions(questions_data, job_description, output_path)
                
                if export_success:
                    results["exported_files"].append({
                        "role": role,
                        "salary_coefficient": salary_coefficient,
                        "file_path": output_path
                    })
                else:
                    results["failed_exports"].append({
                        "role": role,
                        "salary_coefficient": salary_coefficient,
                        "error": "Export işlemi başarısız"
                    })
                    
            except Exception as e:
                logger.error(f"Çoklu export hatası: {e}")
                results["failed_exports"].append({
                    "role": questions_data.get("role", "unknown"),
                    "error": str(e)
                })
        
        # Genel başarı durumu
        if results["failed_exports"]:
            results["success"] = False
        
        logger.info(f"Çoklu export tamamlandı: {len(results['exported_files'])} başarılı, {len(results['failed_exports'])} başarısız")
        return results